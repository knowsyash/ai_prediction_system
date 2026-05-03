from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "AI_Powered_Predictive_Analysis_of_Smartphone_Market_Trends_IEEE_Verified.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["sz", "val", "color", "space"]:
                if key in kwargs[edge]:
                    element.set(qn("w:{}".format(key)), str(kwargs[edge][key]))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Pt(widths[i] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": "BFBFBF"},
                bottom={"val": "single", "sz": "4", "color": "BFBFBF"},
                left={"val": "single", "sz": "4", "color": "BFBFBF"},
                right={"val": "single", "sz": "4", "color": "BFBFBF"},
            )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, txt, fld_end])


def add_paragraph(doc, text="", style=None, align=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_figure_placeholder(doc, number, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "6")
        node.set(qn("w:color"), "A6A6A6")
        borders.append(node)
    p_pr.append(borders)
    r = p.add_run(f"[Insert Figure {number} here: {caption}]")
    r.italic = True
    r.font.color.rgb = RGBColor(80, 80, 80)
    r.font.size = Pt(10)
    p2 = doc.add_paragraph(f"Figure {number}: {caption}")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p2.runs:
        run.font.size = Pt(9)
        run.italic = True
    return p


def add_table(doc, title, headers, rows, widths):
    cap = doc.add_paragraph(title)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.bold = True
        run.font.size = Pt(9)
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_shading(cell, "EDEDED")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8.5)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(8.5)
    set_table_width(table, widths)
    return table


def style_doc(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.line_spacing = 1.0
    styles["Normal"].paragraph_format.space_after = Pt(4)
    for style_name, size in [("Heading 1", 12), ("Heading 2", 10.5), ("Heading 3", 10)]:
        st = styles[style_name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(4)
    styles["Title"].font.name = "Times New Roman"
    styles["Title"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Title"].font.size = Pt(18)
    styles["Title"].font.bold = True


def main():
    doc = Document()
    style_doc(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    header = section.header.paragraphs[0]
    header.text = "IEEE-Style Research Paper | AI-Powered Smartphone Market Trend Prediction"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(90, 90, 90)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("AI-Powered Predictive Analysis of Smartphone Market Trends Using Sentiment-Aware Ensemble Learning")

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = author.add_run(
        "Yash Singhal, Yash Khandelwal, Megha Kundi\n"
        "Department of Computer Science and Engineering\n"
        "Jaypee University of Engineering and Technology, Guna, India\n"
        "231b405@juetguna.in, 231b401@juetguna.in, 231b187@juetguna.in"
    )
    r.font.size = Pt(10)

    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "Abstract", "Keywords", "I. Introduction", "II. Literature Review",
        "III. Methodology / Proposed System", "IV. Implementation",
        "V. Results and Discussion", "VI. Advantages and Limitations",
        "VII. Future Scope", "VIII. Conclusion", "References"
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.2)

    add_heading(doc, "Abstract", 1)
    add_paragraph(doc, (
        "This paper presents an end-to-end AI-powered prediction system for analyzing and forecasting smartphone market "
        "trends from Flipkart customer reviews. The system was designed to solve a practical problem: e-commerce reviews "
        "contain valuable consumer sentiment, but the data is noisy, unevenly distributed, difficult to scrape continuously, "
        "and rarely converted into future-facing product intelligence. The proposed pipeline collects reviews for iPhone 15, "
        "iPhone 16, and iQOO Z10 using a fault-tolerant scraper with user-agent rotation, retry handling, checkpoint saving, "
        "backup CSV generation, and progress logs that allow scraping to resume after interruption or blocking. The collected "
        "data is cleaned through a conservative reviewer-name removal algorithm before sentiment modeling. Three complementary "
        "sentiment approaches are then evaluated: VADER for fast rule-based polarity, BERT for contextual language understanding, "
        "and a domain-specific word-cloud keyword proxy for interpretability. Their outputs are combined through hard voting, "
        "weighted soft voting, and bootstrap bagging to reduce single-model bias. Finally, the best sentiment signal is fused "
        "with linear trend projection and exponential moving average forecasting to predict next-month ratings. Experiments on "
        "4,021 reviews show that BERT achieves 86.99% individual accuracy, soft voting improves balanced performance with "
        "0.5424 macro-F1 and 0.5916 Pearson correlation, and bagging reduces RMSE to 0.8973. The final system predicts ratings "
        "of 4.477, 4.467, and 4.339 for iPhone 15, iPhone 16, and iQOO Z10 respectively, demonstrating that sentiment-aware "
        "ensemble forecasting can transform raw public reviews into reliable and interpretable market trend intelligence."
    ))

    add_heading(doc, "Keywords", 1)
    add_paragraph(doc, "Sentiment analysis, smartphone reviews, ensemble learning, BERT, VADER, market forecasting, exponential smoothing.")

    add_heading(doc, "I. Introduction", 1)
    add_heading(doc, "A. Background", 2)
    add_paragraph(doc, (
        "Smartphone markets move quickly because consumer expectations around camera quality, battery endurance, "
        "display performance, software stability, price, and after-sales experience change rapidly after launch. "
        "Public review platforms therefore contain a dense record of post-purchase perception. Unlike structured "
        "sales dashboards, review text captures why users reward or reject a device, including subtle concerns such "
        "as heating, charging behavior, network reception, and perceived value. Sentiment analysis and opinion mining "
        "have long been used to organize such subjective text [1], [2], but market-facing prediction requires an "
        "additional step: converting noisy and delayed consumer opinions into a forward-looking quantitative signal."
    ))
    add_paragraph(doc, (
        "The project studied in this paper addresses that gap through a complete pipeline. It performs web scraping, "
        "data cleaning, sentiment modeling, ensemble aggregation, and next-month rating prediction. The central "
        "research premise is that recent review sentiment, when combined with historical rating trends, can act as "
        "a leading indicator of product trajectory."
    ))
    add_heading(doc, "B. Problem Statement", 2)
    add_paragraph(doc, (
        "Existing review analytics workflows often rely on a single sentiment method or report only aggregate "
        "positive-negative ratios. Such summaries are useful for retrospective monitoring but inadequate for "
        "forecasting. Lexicon models are fast but may miss product-specific context; transformer models are accurate "
        "but computationally heavy and can be biased toward dominant classes; keyword models are interpretable but "
        "lose syntax and negation. A robust smartphone trend prediction system must therefore reconcile model diversity, "
        "class imbalance, noisy text, and short historical time series."
    ))
    add_paragraph(doc, (
        "A second practical problem emerged during data collection. Public e-commerce pages are not designed as stable "
        "research APIs. During scraping, requests may fail because of temporary network issues, anti-bot behavior, repeated "
        "access from the same IP address, changed HTML classes, empty pages, or rate-limiting responses. Restarting a scraper "
        "from page one after every interruption wastes time, increases duplicate traffic, and can further increase blocking "
        "risk. Therefore, this project required not only a sentiment model but also a resilient data acquisition method that "
        "could pause, save, recover, and resume from the last known successful point."
    ))
    add_heading(doc, "C. Dataset Scope and Source Rationale", 2)
    add_paragraph(doc, (
        "The study intentionally uses Flipkart review data rather than mixing reviews from multiple sources. Flipkart was "
        "selected because it is a major Indian e-commerce marketplace, the target products are actively reviewed there, and "
        "review pages provide the minimum fields required for supervised evaluation: star rating, review title, review body, "
        "and review date. The star rating is essential because it acts as the ground-truth label for positive, neutral, and "
        "negative sentiment mapping. Without that rating field, the project would require manual labeling or weak labels, "
        "which would reduce reproducibility."
    ))
    add_paragraph(doc, (
        "Other sources such as Amazon, YouTube comments, social media posts, and general web forums were deliberately excluded "
        "from the first version. They have different schemas, different moderation policies, different levels of purchase "
        "verification, and different kinds of noise. Combining them too early would create a cross-platform normalization "
        "problem: a five-star product review, a short social media complaint, and a video comment do not carry the same signal. "
        "Using one platform kept the dataset internally consistent and allowed the project to focus on the core research "
        "question: whether cleaned review sentiment can improve next-month product rating forecasts."
    ))
    add_heading(doc, "D. Objectives", 2)
    add_paragraph(doc, (
        "The objectives of this work are to design a reproducible pipeline for smartphone review collection and "
        "cleaning, compare multiple sentiment models against star-rating ground truth, construct ensemble models that "
        "balance accuracy and rating-scale error, and produce next-month product rating forecasts through a hybrid "
        "combination of sentiment and time-series signals."
    ))

    add_heading(doc, "II. Literature Review", 1)
    add_paragraph(doc, (
        "Pang, Lee, and Vaithyanathan [3] established early supervised sentiment classification by applying machine "
        "learning methods to review polarity, showing that sentiment classification is more difficult than topic "
        "classification because subjective language is context-sensitive. Turney [4] demonstrated an unsupervised "
        "semantic orientation approach, using adjective and adverb phrases to classify reviews as recommended or not "
        "recommended. These works provide the conceptual basis for treating user reviews as predictive signals rather "
        "than simple text records."
    ))
    add_paragraph(doc, (
        "Hu and Liu [5] moved sentiment analysis closer to product analytics by mining and summarizing customer reviews "
        "around product features. Their feature-based summarization is especially relevant to smartphones, where the "
        "same overall rating may arise from different factors such as camera quality, battery life, or delivery experience. "
        "Liu's broader survey [1] further systematized opinion mining tasks, including sentiment classification, target "
        "identification, and opinion summarization."
    ))
    add_paragraph(doc, (
        "Hutto and Gilbert [6] introduced VADER, a parsimonious rule-based sentiment model designed for short informal "
        "text. Its practical strength lies in speed and human-readable scoring rules for punctuation, capitalization, "
        "degree modifiers, and negation. However, lexicon-based systems are limited when a review contains domain-specific "
        "phrases or indirect sentiment. Transformer models address this limitation. Vaswani et al. [7] introduced the "
        "self-attention architecture that enabled modern contextual language models, while Devlin et al. [8] proposed "
        "BERT, which learns bidirectional contextual representations and has become a strong baseline for text classification."
    ))
    add_paragraph(doc, (
        "Forecasting literature provides the second foundation of this work. Gardner [9] reviewed exponential smoothing "
        "as a practical forecasting family that reacts to recent changes while retaining historical memory. Linear trend "
        "models remain interpretable and useful when long-term rating movement is approximately monotonic. Ensemble "
        "learning supplies the third foundation: Breiman's bagging [10] reduces model variance through bootstrap aggregation, "
        "and Friedman [11] showed how additive model combinations can improve predictive performance. In review analytics, "
        "these ideas motivate combining sentiment classifiers rather than selecting a single winner. Compared with the "
        "above works, the present system contributes an end-to-end implementation that joins review scraping, transformer "
        "and lexicon sentiment, ensemble correction, and next-month smartphone rating prediction in one reproducible pipeline."
    ))

    add_heading(doc, "III. Methodology / Proposed System", 1)
    add_heading(doc, "A. System Architecture", 2)
    add_paragraph(doc, (
        "The proposed architecture follows a staged data-to-forecast design. The first layer acquires reviews from "
        "Flipkart product pages and stores rating, review text, reviewer name, and date fields in CSV format. The second "
        "layer performs cleaning, including conservative reviewer-name removal from the end of review strings. The third "
        "layer estimates sentiment through VADER, BERT, and a TF-IDF word-cloud proxy. The fourth layer aggregates model "
        "outputs through voting and bootstrap-based ensembles. The final layer combines the best sentiment signal with "
        "monthly rating trends to forecast the next-month average rating."
    ))
    add_paragraph(doc, (
        "The project is organized as a reproducible seven-stage workspace. The folder `1_data_scrapping` stores separate "
        "Flipkart scrapers for iPhone 15, iPhone 16, and iQOO Z10. The folder `2_dataset_final_folder` contains both "
        "pre-cleaning backups and final cleaned datasets. The folder `3_sentimental_analysis` contains the comparative "
        "sentiment experiment, while `4_enemble_analysis` stores ensemble outputs, reliability plots, cumulative ratings, "
        "and model summaries. The folder `5_next_month` contains the forecasting script and final prediction artifacts. "
        "The folder `6_dashboard` turns the CSV and chart outputs into an interactive Dash dashboard, and `docs` records "
        "the technical reasoning behind data cleaning, sentiment analysis, ensembling, forecasting, and the mathematical "
        "walkthrough."
    ))
    add_table(
        doc,
        "Table I. Project Artifact Map",
        ["Stage", "Folder / File", "Main Role", "Output Produced"],
        [
            ["1", "1_data_scrapping", "Flipkart review extraction using Requests, BeautifulSoup, rotating user agents, retry logic, and checkpoint saving.", "Raw review CSV files for each smartphone."],
            ["2", "2_dataset_final_folder", "Reviewer-name removal, text normalization, backup preservation, and final data preparation.", "Cleaned CSVs used by all later scripts."],
            ["3", "3_sentimental_analysis/YS", "Comparison of VADER, BERT, and TF-IDF word-cloud sentiment models.", "comparison_results_v2.csv, comparison_summary_v2.csv, confusion matrices."],
            ["4", "4_enemble_analysis", "Hard voting, soft voting, and bagging over base sentiment predictions.", "ensemble_results_v2.csv, ensemble_summary_v2.csv, reliability and RMSE plots."],
            ["5", "5_next_month", "Hybrid forecasting from linear trend, exponential smoothing, and soft-voting sentiment.", "next_month_prediction_v2.csv and trend charts."],
            ["6", "6_dashboard", "Interactive visualization of model metrics, forecast components, and time-series behavior.", "Dash web dashboard."],
            ["7", "docs", "Human-readable project documentation and mathematical walkthrough.", "Design reports and implementation explanation."],
        ],
        [700, 1800, 3900, 2600],
    )
    add_figure_placeholder(doc, 1, "System Architecture")

    add_heading(doc, "B. Algorithms and Models Used", 2)
    add_paragraph(doc, (
        "VADER computes a compound polarity score using a sentiment lexicon and rule-based intensity modifiers. The "
        "BERT component uses a pretrained transformer sentiment classifier that maps review text into contextual token "
        "representations and returns class probabilities. The word-cloud category proxy derives positive and negative "
        "keyword weights from TF-IDF scores, producing an interpretable domain-specific baseline. Hard voting uses the "
        "majority sentiment label across base models; soft voting averages normalized score signals; and bagging estimates "
        "rating-level sentiment through repeated bootstrap aggregation. For forecasting, the system uses linear regression "
        "on monthly mean ratings, single exponential smoothing, and a recent soft-voting sentiment signal."
    ))
    add_paragraph(doc, (
        "The final prediction is computed as Final Rating = w1L + w2E + w3S, where L is the linear trend forecast, E is "
        "the exponential smoothing forecast, and S is the recent sentiment forecast. The weights are dynamic. Devices with "
        "six or more historical months use 0.40, 0.35, and 0.25 for linear trend, smoothing, and sentiment respectively; "
        "devices with three to five months use 0.30, 0.35, and 0.35; devices with fewer than three months depend more "
        "heavily on recent sentiment."
    ))
    add_paragraph(doc, (
        "The scraper layer uses several safeguards because e-commerce review pages are semi-structured and can change over "
        "time. Each request is made through a persistent session for connection reuse. Browser headers are rotated to reduce "
        "blocking, and randomized delays are inserted between requests. A page is parsed by scanning likely review-card "
        "containers and then falling back to generic div-pattern extraction. Noise terms such as certified buyer labels, "
        "report-abuse links, storage/color markers, and pagination text are filtered before a row is accepted. Deduplication "
        "is performed using a composite key built from rating, title prefix, and review-text prefix. The iPhone 16 and iQOO "
        "Z10 scrapers additionally extract or infer review month values and implement resume logic through progress files."
    ))
    add_paragraph(doc, (
        "The need for this design came from practical scraping failures. When many pages are requested sequentially, the host "
        "may stop returning normal review cards, the network may time out, or the scraper may be interrupted manually. The "
        "scripts therefore avoid treating scraping as one long all-or-nothing operation. They save periodically after a small "
        "batch of pages, maintain a `simple_save.txt` log with timestamps and page-level counts, preserve backup CSV files "
        "before rewriting, and track the last successful page. The iQOO Z10 scraper additionally writes `last_page.txt` with "
        "a parseable `LAST_SUCCESSFUL_PAGE` value. This means that after a stop, the scraper can continue from the next page "
        "or the nearest safe checkpoint instead of starting again from the first page. In practical terms, this is what made "
        "it possible to keep collecting more reviews even when IP-based blocking, empty responses, or runtime interruptions "
        "occurred."
    ))
    add_paragraph(doc, (
        "The checkpoint size is intentionally small. In the code, the scraper saves every five pages rather than waiting until "
        "the end of the product crawl. Depending on the number of reviews extracted per page, this behaves like a small review "
        "batch checkpoint. The project description often refers to this idea as saving after a group of reviews; the exact "
        "implementation is page-based because page count is easier to track reliably during crawling. Each checkpoint writes "
        "new rows to CSV, removes duplicates, clears in-memory rows, and logs the current state. This reduces memory usage, "
        "protects already collected reviews, and prevents duplicate records after resuming."
    ))
    add_paragraph(doc, (
        "The cleaning layer is intentionally conservative. The `clean_names.py` script does not simply remove the last two "
        "words of every review. Instead, it checks whether the trailing two- or three-token block looks like a reviewer name. "
        "A candidate name token must be alphabetic, lowercase after normalization, between 2 and 20 characters, and absent "
        "from a large English-word blocklist. For two-token names, the preceding token must be either punctuation-terminated "
        "or a recognized content word. For three-token names, the preceding token must end with punctuation. This design "
        "protects valid endings such as 'very good', 'this phone', and 'best part' while removing examples such as 'sarath "
        "kumar' or 'bijaya mohanty'."
    ))
    add_table(
        doc,
        "Table II. Dataset Preparation and Current Cleaned Snapshot",
        ["Device", "Rows Before Cleaning", "Rows After Current Cleaning", "Text Rows Modified", "Current Rating Distribution"],
        [
            ["iPhone 15", "3315", "3315", "456", "1:143, 2:49, 3:105, 4:414, 5:2604"],
            ["iPhone 16", "719", "719", "99", "1:28, 2:7, 3:22, 4:86, 5:576"],
            ["iQOO Z10", "57", "46", "46", "1:4, 2:3, 3:1, 4:8, 5:30"],
        ],
        [1300, 1500, 1500, 1500, 3200],
    )
    add_heading(doc, "C. Workflow Explanation", 2)
    add_paragraph(doc, (
        "The workflow begins by extracting raw reviews and converting them into structured tabular records. After cleaning, "
        "each review is mapped to a ground-truth sentiment label using its star rating: one or two stars are negative, "
        "three stars are neutral, and four or five stars are positive. The three sentiment methods are evaluated using "
        "accuracy, macro-F1, RMSE, Pearson correlation, and Spearman correlation. Ensemble outputs are then measured using "
        "the same metrics to identify the best signal for forecasting. Finally, review dates are grouped by month to form "
        "time series for each device, and the three forecasting signals are combined into the next-month rating prediction."
    ))
    add_paragraph(doc, (
        "The paper distinguishes between the current cleaned dataset snapshot and the archived model-output dataset used for "
        "the reported experiment. The saved comparison and ensemble result files contain 4,021 usable review rows: 3,245 for "
        "iPhone 15, 719 for iPhone 16, and 57 for iQOO Z10. The current cleaned CSV snapshot contains a later iQOO Z10 file "
        "with 46 rows after additional cleaning. Therefore, the performance metrics, plots, cumulative ratings, and forecast "
        "tables are cited from the archived experiment artifacts so that the reported values match `comparison_results_v2.csv`, "
        "`ensemble_results_v2.csv`, `ensemble_summary_v2.csv`, `cumulative_ratings_v2.csv`, and `next_month_prediction_v2.csv`."
    ))
    add_paragraph(doc, (
        "At the individual-review level, the system stores both discrete sentiment labels and continuous predicted ratings. "
        "For VADER, the compound polarity score is linearly rescaled from [-1, 1] to [1, 5] by the transformation rating = "
        "(compound x 2) + 3. For BERT, the implemented DistilBERT SST-2 pipeline returns a positive/negative sentiment "
        "decision; the script converts this output into a positive probability and maps it to the rating scale using rating = "
        "1 + 4 x P(positive). Values near 1 indicate strongly negative text, while values near 5 indicate strongly positive "
        "text. For the word-cloud proxy, the model starts from a neutral base score and adjusts the value upward or downward "
        "according to domain keywords learned from positive and negative review groups."
    ))
    add_paragraph(doc, (
        "The choice of the three sentiment models was also deliberate. VADER was included because it is fast, transparent, "
        "and robust for short informal text, making it suitable for quick baselines and large batches. BERT was included "
        "because transformer encoders understand context, contrast, and indirect sentiment better than word lists. The "
        "word-cloud proxy was included because it is domain-specific and interpretable: it makes visible which product-review "
        "terms are contributing to the score. These three models have different failure modes, which is precisely why they "
        "are useful together."
    ))
    add_paragraph(doc, (
        "The ensemble stage then resolves disagreement between models. Hard voting converts each continuous prediction into "
        "a class and selects the majority label. Soft voting preserves numeric information by averaging the continuous "
        "ratings, which prevents a score such as 3.95 from being treated the same as a clearly neutral 3.00. Bagging creates "
        "50 bootstrap samples with replacement and averages repeated soft-voting estimates, reducing sensitivity to unusual "
        "reviews or small sample perturbations. The forecasting stage selects soft voting as the primary sentiment signal "
        "because it offers the best balance of macro-F1 and correlation with actual ratings."
    ))
    add_paragraph(doc, (
        "Soft voting in the implementation is not an unweighted average. When BERT outputs are available, the script combines "
        "the rating-like scores using 0.40 weight for VADER, 0.40 for BERT, and 0.20 for the word-cloud proxy. This preserves "
        "the strength of the two more general sentiment engines while still retaining a contribution from the domain keyword "
        "model. If BERT is unavailable, the fallback uses 0.60 VADER and 0.40 word-cloud weighting. The soft-vote score is "
        "then mapped back to sentiment classes using rating thresholds: scores at or above 3.75 are positive, scores at or "
        "below 2.25 are negative, and the middle region is neutral."
    ))
    add_paragraph(doc, (
        "Bagging is used to examine stability. The script creates 50 bootstrap samples, each containing approximately 80% of "
        "the dataset sampled with replacement. Since the base models are deterministic after their predictions are generated, "
        "bagging does not retrain a new neural network; instead, it measures how stable the ensemble rating estimate remains "
        "under resampled review populations. The resulting `bagging_pred_std` field records per-review uncertainty. Lower "
        "standard deviation indicates that the prediction is less sensitive to which reviews happened to be sampled."
    ))

    add_heading(doc, "IV. Implementation", 1)
    add_heading(doc, "A. Tools and Technologies", 2)
    add_paragraph(doc, (
        "The system is implemented in Python. Pandas and NumPy support dataset manipulation, monthly aggregation, and "
        "numeric computation. Scikit-learn and SciPy are used for classification metrics, RMSE, and correlation analysis. "
        "The vaderSentiment package provides VADER scoring, while Hugging Face Transformers and PyTorch support the BERT "
        "sentiment model. Matplotlib, Seaborn, Plotly, Dash, and wordcloud support visualization, dashboarding, and "
        "interpretable keyword analysis."
    ))
    add_heading(doc, "B. Step-by-Step Working", 2)
    add_paragraph(doc, (
        "First, the scraper collects review pages using browser-like request headers, throttling, retry logic, and progress "
        "logging. Second, cleaning scripts preserve the review content while removing trailing reviewer names that would "
        "otherwise distort sentiment tokens. Third, the sentiment comparison script processes the cleaned data and exports "
        "per-review predictions and aggregate metrics. Fourth, the ensemble script imports those predictions, computes hard "
        "voting, soft voting, and bagging outputs, and generates reliability plots. Fifth, the forecasting script reads the "
        "cleaned product files and ensemble results, derives monthly signals, applies dynamic weights, and exports forecast "
        "tables and charts. The dashboard folder then serves the generated summaries for interactive inspection."
    ))
    add_paragraph(doc, (
        "The data acquisition scripts differ slightly by device because the available Flipkart page structures and review "
        "metadata vary. The iPhone 15 scraper focuses on rating, title, and review text extraction and saves progress every "
        "five pages. The iPhone 16 scraper extends this logic by extracting review dates from explicit month-year strings "
        "or relative expressions such as months ago, and it also attempts city extraction before removing location text from "
        "the final review content. The iQOO Z10 scraper implements a two-level parsing strategy: it first searches common "
        "Flipkart review-card classes and data-review-id attributes, then falls back to generic text segmentation if the "
        "structured parse fails. It also records the last successful page in `last_page.txt` and estimates a resume page "
        "from existing CSV length."
    ))
    add_paragraph(doc, (
        "During validation, each candidate review must satisfy length, noise, and uniqueness constraints. Titles shorter than "
        "four characters, numeric-only titles, page navigation fragments, storage/color labels, and platform boilerplate are "
        "discarded. Review text is limited to a safe length range so that extremely short fragments and large page-level text "
        "blocks are not treated as genuine reviews. Before each CSV write, the script creates a backup where applicable, "
        "appends new rows, removes duplicates on title and review text, and writes the consolidated dataset back to disk."
    ))
    add_paragraph(doc, (
        "The progress-management design is one of the most important implementation choices. The scraper records every major "
        "event with a timestamp: starting page, pages with successful review counts, pages with no reviews, network retry "
        "messages, checkpoint saves, interruption events, and final totals. This text log is not only for debugging; it is a "
        "recovery record. If the program stops because the internet connection fails, the website blocks repeated requests, "
        "or the user interrupts execution, the log shows the last page that produced usable reviews. The iPhone 16 scraper "
        "parses this progress file to calculate a resume page based on the nearest previous multiple-of-five checkpoint. The "
        "iQOO Z10 scraper combines CSV length estimation with `last_page.txt`, choosing the safer higher page number before "
        "continuing. This avoided repeated scraping from the beginning and reduced load on the source website."
    ))
    add_paragraph(doc, (
        "The project also had to handle changing page structures. Flipkart review cards do not always expose clean semantic "
        "selectors, and different products may render review text differently. The iQOO Z10 scraper therefore tries structured "
        "selectors such as review-card attributes first, then falls back to a generic parse that separates page text by delimiters "
        "and validates rating-title-review patterns. The iPhone scrapers scan div text line-by-line and reject blocks containing "
        "page-level summary phrases such as ratings and reviews or user reviews sorted. This hybrid parsing strategy is less "
        "elegant than an official API, but it is practical for real scraped web data."
    ))
    add_heading(doc, "E. Why EMA and Weighted Forecasting Were Used", 2)
    add_paragraph(doc, (
        "Exponential moving average, implemented as single exponential smoothing, was chosen because smartphone review streams "
        "are time-dependent but not large enough for heavy time-series models. A product can receive sudden negative sentiment "
        "after a software update, price change, delivery issue, or competitor launch. A simple historical mean would react too "
        "slowly, while a model that only uses the latest month would overreact to temporary noise. EMA solves this by blending "
        "the latest observation with the previous smoothed value. The smoothing factor alpha controls responsiveness: this "
        "project uses alpha = 0.30 when at least eight monthly observations exist and alpha = 0.45 when the history is shorter. "
        "Thus, mature devices receive steadier forecasts, while newer devices respond more strongly to recent movement."
    ))
    add_paragraph(doc, (
        "The final forecast is weighted instead of purely model-driven because each signal answers a different question. Linear "
        "trend answers whether the product has been gradually improving or declining. EMA answers what the recent momentum "
        "looks like. Sentiment answers what the latest review text says beyond the star rating. If a device has six or more "
        "months of history, the project gives more weight to trend and smoothing: 0.40 linear, 0.35 EMA, and 0.25 sentiment. "
        "If the device has three to five months, the weights are 0.30, 0.35, and 0.35 because recent sentiment still matters "
        "strongly. With fewer than three months, sentiment becomes the dominant signal because the time series is too short "
        "to support a reliable trend."
    ))
    add_paragraph(doc, (
        "The sentiment comparison script evaluates each method at two levels. At the review level, it stores text, ground-truth "
        "rating, predicted sentiment label, and predicted rating-like score. At the summary level, it computes accuracy, "
        "macro-F1, RMSE, Pearson correlation, and Spearman correlation. Accuracy is retained for readability, but macro-F1 "
        "is emphasized because the dataset contains many more positive reviews than neutral or negative reviews. RMSE is "
        "included because the final task is rating prediction rather than only class prediction."
    ))
    add_table(
        doc,
        "Table III. Key Mathematical Operations Used in the Pipeline",
        ["Operation", "Formula / Rule", "Purpose in Project"],
        [
            ["VADER rating conversion", "rating = (compound x 2) + 3", "Maps polarity from [-1, 1] to the 1-5 product-rating scale."],
            ["BERT rating conversion", "rating = 1 + 4 x P(positive)", "Maps the DistilBERT SST-2 positive probability to the 1-5 rating scale."],
            ["Hard voting", "class = majority(base model labels)", "Produces an interpretable class decision from three model votes."],
            ["Soft voting", "rating = 0.40 VADER + 0.40 BERT + 0.20 Word Cloud", "Preserves confidence while weighting the two general sentiment engines more strongly."],
            ["Bagging", "50 bootstrap samples; sample fraction = 0.80", "Measures stability and reduces sensitivity to resampled review populations."],
            ["Linear trend", "y = mx + c", "Projects long-term monthly rating movement into the target month."],
            ["Exponential smoothing", "S_t = alpha y_t + (1-alpha)S_(t-1)", "Captures recent momentum without ignoring historical behavior."],
            ["Final forecast", "F = w1L + w2E + w3S", "Combines trend, smoothing, and sentiment into the next-month rating."],
        ],
        [1900, 3300, 3800],
    )
    add_heading(doc, "C. Dashboard Implementation", 2)
    add_paragraph(doc, (
        "The dashboard is implemented with Dash and Plotly. It loads `comparison_summary_v2.csv`, `ensemble_summary_v2.csv`, "
        "`ensemble_results_v2.csv`, and `next_month_prediction_v2.csv` directly from the project folders, then rebuilds "
        "monthly time-series values from the cleaned product datasets. A month parser converts strings such as Jan 2026 into "
        "timestamps. The first dashboard view displays the final forecast breakdown for a selected device, showing each "
        "component value, its assigned weight, and its contribution to the final prediction. The second view overlays actual "
        "monthly ratings, soft-voting monthly predictions, and the next-month forecast with confidence interval bounds. This "
        "makes the dashboard a verification surface, not only a presentation layer."
    ))
    add_heading(doc, "D. Reproducibility Notes", 2)
    add_paragraph(doc, (
        "All major outputs are stored as CSV or PNG artifacts, which makes the pipeline auditable. Intermediate files such as "
        "`comparison_results_v2.csv` and `ensemble_results_v2.csv` preserve per-review predictions, allowing any aggregate "
        "metric to be recomputed. Summary files provide compact tables for reporting, while generated plots support visual "
        "inspection of confusion matrices, metric comparisons, reliability behavior, cumulative ratings, trend lines, and "
        "component contributions. The separation of raw, cleaned, model, ensemble, forecast, and dashboard folders also makes "
        "it possible to rerun one stage without overwriting the full project."
    ))
    add_paragraph(doc, (
        "The project outputs are intentionally redundant. Per-review result files preserve the detailed prediction trail, "
        "summary CSVs preserve compact metrics, and PNG plots preserve visual evidence. This is useful in academic reporting "
        "because a table alone does not reveal class imbalance or reliability behavior. Confusion matrices show which classes "
        "are being confused; correlation plots show whether predicted ratings move with actual ratings; reliability plots show "
        "whether bootstrap predictions are stable; and component-breakdown plots show how much each forecast signal contributes "
        "to the final value."
    ))
    add_figure_placeholder(doc, 2, "Implementation Flow")

    add_heading(doc, "V. Results and Discussion", 1)
    add_heading(doc, "A. Performance Analysis", 2)
    add_paragraph(doc, (
        "The experimental dataset contains 4,021 reviews: 3,245 for iPhone 15, 719 for iPhone 16, and 57 for iQOO Z10. "
        "The review distribution is strongly positive, with actual positive-rate values of 91.2%, 92.1%, and 86.0% "
        "respectively. This imbalance explains why accuracy alone is insufficient. BERT achieves the strongest single-model "
        "accuracy, but macro-F1 and RMSE reveal that simpler and ensemble methods handle minority classes and rating-scale "
        "error more effectively."
    ))
    add_paragraph(doc, (
        "The sentiment comparison produces an important practical insight. BERT reaches 0.8699 accuracy overall, which "
        "confirms the value of contextual transformer representations for common review polarity. However, its macro-F1 of "
        "0.4758 is lower than the word-cloud proxy's 0.5094 and VADER's 0.5000, indicating that the high accuracy is partly "
        "driven by the dominant positive class. The word-cloud proxy obtains the lowest single-model RMSE of 1.1245 because "
        "domain-specific keywords such as camera, battery, amazing, bad, heating, and value align closely with rating-scale "
        "movement. This explains why the project did not simply adopt BERT as the final model."
    ))
    add_table(
        doc,
        "Table IV. Overall Sentiment Model Comparison",
        ["Method", "N", "Accuracy", "Macro-F1", "RMSE", "Pearson r", "Spearman rho"],
        [
            ["VADER", "4021", "0.7690", "0.5000", "1.1606", "0.4529", "0.3018"],
            ["BERT", "4021", "0.8699", "0.4758", "1.2094", "0.5423", "0.3520"],
            ["Word Cloud", "4021", "0.7120", "0.5094", "1.1245", "0.4798", "0.3219"],
            ["Hard Voting", "4021", "0.7819", "0.5389", "1.0533", "0.5523", "0.3703"],
            ["Soft Voting", "4021", "0.8386", "0.5424", "0.9599", "0.5916", "0.3565"],
            ["Bagging", "4021", "0.8374", "0.3776", "0.8973", "0.5900", "0.3560"],
        ],
        [1800, 900, 1200, 1200, 1100, 1300, 1500],
    )
    add_paragraph(doc, (
        "Soft voting provides the most balanced sentiment classifier, achieving the highest macro-F1 and Pearson correlation. "
        "Bagging achieves the lowest RMSE, indicating the best rating-scale proximity. The result supports the central design "
        "choice of preserving model diversity rather than selecting a single sentiment classifier."
    ))
    add_paragraph(doc, (
        "The ensemble results show that hard voting improves macro-F1 to 0.5389 by reducing isolated model mistakes, but it "
        "still discards score intensity. Soft voting improves Pearson correlation to 0.5916 and macro-F1 to 0.5424 because "
        "it uses the continuous strength of each base model's prediction. Bagging produces the lowest RMSE of 0.8973, which "
        "means its predictions are closest to actual star ratings on average. The final system therefore treats soft voting "
        "as the most balanced live sentiment signal and bagging as evidence that bootstrap aggregation can stabilize the "
        "rating-scale estimate."
    ))
    add_table(
        doc,
        "Table V. Next-Month Forecast Results",
        ["Device", "Target Month", "History", "Last Rating", "Linear", "Exp. Smooth", "Sentiment", "Predicted"],
        [
            ["iPhone 16", "Feb 2026", "14 months", "4.455", "4.530", "4.615", "4.160", "4.467"],
            ["iPhone 15", "Feb 2025", "17 months", "4.633", "4.566", "4.618", "4.139", "4.477"],
            ["iQOO Z10", "Mar 2026", "3 months", "4.643", "4.606", "4.445", "4.005", "4.339"],
        ],
        [1500, 1300, 1200, 1100, 1000, 1300, 1200, 1100],
    )
    add_paragraph(doc, (
        "The forecast values remain positive for all products but reveal meaningful differences. iPhone 15 and iPhone 16 "
        "show stable high-rating trajectories because their long historical records support stronger trend and smoothing "
        "weights. iQOO Z10, with only three months of history, is more sensitive to its recent sentiment signal and therefore "
        "receives a lower prediction despite a strong last observed rating."
    ))
    add_table(
        doc,
        "Table VI. Device-Level Cumulative Rating Summary",
        ["Device", "Reviews", "Actual Mean", "Actual Positive %", "VADER Mean", "BERT Mean", "Word-Cloud Mean"],
        [
            ["iPhone 16", "719", "4.634", "92.1", "3.942", "4.247", "4.421"],
            ["iPhone 15", "3245", "4.598", "91.2", "3.840", "4.398", "4.219"],
            ["iQOO Z10", "57", "4.316", "86.0", "4.057", "3.808", "4.296"],
        ],
        [1500, 1000, 1300, 1500, 1300, 1300, 1500],
    )
    add_paragraph(doc, (
        "The device-level table highlights another design issue: mean rating and positive percentage are high for every "
        "device, but the models estimate different absolute rating levels. VADER is conservative for the iPhone devices, "
        "BERT is relatively strong for iPhone 15 and iPhone 16, and the word-cloud proxy closely tracks iQOO Z10 despite "
        "the smaller sample. These differences justify ensemble construction because each model captures a different view "
        "of the same consumer-opinion signal."
    ))
    add_figure_placeholder(doc, 3, "Output Results")

    add_heading(doc, "B. Discussion", 2)
    add_paragraph(doc, (
        "The findings show that market sentiment forecasting benefits from a separation between interpretation and prediction. "
        "BERT is highly effective at broad polarity recognition, but the ensemble methods are more useful when the output must "
        "behave like a rating-scale forecast. The system's dynamic weighting strategy also reduces overreaction: mature products "
        "are guided more by history, while newer products are guided more by recent opinion. This design is particularly suitable "
        "for e-commerce markets where review volume, launch recency, and class imbalance differ across products."
    ))
    add_paragraph(doc, (
        "The observed forecasting behavior also matches the design intention. The iPhone 15 and iPhone 16 have long histories, "
        "so the final predictions are controlled mainly by trend and EMA. Their recent sentiment signal is lower than their "
        "last observed rating, which prevents the forecast from simply copying the high latest rating. The iQOO Z10 has only "
        "three months of history, so sentiment receives a larger weight. As a result, its forecast drops more sharply, which "
        "is appropriate for a product where a small number of recent reviews can materially change the estimated trajectory."
    ))
    add_paragraph(doc, (
        "From a research perspective, the main result is not only that one method scored better than another. The project shows "
        "that the data pipeline itself is part of the model. If scraping is interrupted and data is duplicated, if names remain "
        "attached to review text, if star ratings are mapped poorly, or if class imbalance is ignored, the final forecast will "
        "be unreliable regardless of the machine-learning algorithm. The system therefore treats data acquisition, cleaning, "
        "sentiment modeling, ensembling, and forecasting as a single chain where every stage must preserve signal quality."
    ))

    add_heading(doc, "VI. Advantages and Limitations", 1)
    add_paragraph(doc, (
        "The proposed system offers several advantages. It is end-to-end, reproducible, and interpretable at both model and "
        "business levels. It compares heterogeneous sentiment models, uses ensemble learning to reduce single-model bias, and "
        "produces concrete next-month rating forecasts instead of only descriptive dashboards. The inclusion of VADER and TF-IDF "
        "baselines also makes the system computationally flexible, while BERT provides contextual understanding where review text "
        "is linguistically complex."
    ))
    add_paragraph(doc, (
        "A major practical advantage is fault tolerance during scraping. Since reviews are saved at checkpoints and backed up "
        "before rewriting, the system does not lose all work when crawling stops. The progress log provides traceability, the "
        "resume mechanism reduces duplicate traffic, and duplicate-key filtering prevents repeated rows from contaminating "
        "the dataset. These engineering details are important because most academic prototypes describe scraping as a single "
        "step, whereas real scraping often fails repeatedly due to IP throttling, unstable HTML, or connection limits."
    ))
    add_paragraph(doc, (
        "Another advantage is methodological balance. VADER supplies speed and interpretability, BERT supplies contextual "
        "understanding, the word-cloud proxy supplies domain visibility, soft voting supplies balanced prediction, bagging "
        "supplies robustness testing, and EMA supplies recent-momentum awareness. The final forecast is therefore not dependent "
        "on one black-box model or one fragile statistic."
    ))
    add_paragraph(doc, (
        "The limitations are equally important. The dataset is sourced from one e-commerce platform, so platform-specific reviewer "
        "behavior may affect generalization. Star ratings are treated as ground truth, although users sometimes provide ratings "
        "that conflict with review text. The iQOO Z10 dataset is much smaller than the iPhone datasets, which increases forecast "
        "uncertainty. The BERT component is computationally heavier than rule-based methods, and the current model does not "
        "explicitly extract product aspects such as camera, battery, display, or service quality."
    ))
    add_paragraph(doc, (
        "The single-platform choice is a limitation but also a controlled experimental decision. Flipkart reviews represent an "
        "Indian e-commerce context and may not generalize to global buyer behavior, offline retail sentiment, or social media "
        "discussion. In addition, scraped pages may omit reviews hidden behind dynamic rendering, pagination changes, or anti-bot "
        "responses. The confidence intervals in the current forecasting script are heuristic rather than fully probabilistic. "
        "Finally, BERT was used as a pretrained sentiment pipeline rather than a smartphone-specific fine-tuned model, so it may "
        "misread domain language such as heating, lag, dynamic island, charging, or value for money."
    ))

    add_heading(doc, "VII. Future Scope", 1)
    add_paragraph(doc, (
        "Future work can extend the system in five directions. First, aspect-based sentiment analysis can be added to separate "
        "camera, battery, display, performance, and service-related sentiment. Second, review data can be collected from multiple "
        "platforms to reduce platform bias. Third, transformer models can be fine-tuned on smartphone-specific review data to "
        "improve minority-class recall. Fourth, probabilistic forecasting can replace fixed confidence intervals with calibrated "
        "uncertainty estimates. Fifth, the dashboard can be expanded into a decision-support system that alerts analysts when "
        "recent sentiment deviates materially from historical trend."
    ))
    add_paragraph(doc, (
        "A future version can also add multi-source ingestion after a stronger normalization layer is built. Amazon, official "
        "brand stores, YouTube reviews, Reddit discussions, and social media posts can be added as separate channels rather "
        "than merged directly. Each channel should have its own trust score, schema, language filter, duplicate-removal logic, "
        "and sentiment calibration. This would allow the system to compare verified-purchase sentiment with broader public "
        "opinion while avoiding the confusion that comes from mixing fundamentally different text sources."
    ))
    add_paragraph(doc, (
        "The scraping layer can be improved by adding proxy rotation, adaptive delay scheduling, browser automation for pages "
        "that render reviews dynamically, and explicit HTTP-status logging. The modeling layer can be improved by fine-tuning "
        "a transformer on the collected smartphone reviews and by training an aspect-based classifier. The forecasting layer "
        "can be improved by backtesting predictions month by month, estimating calibrated prediction intervals, and comparing "
        "EMA against ARIMA, Prophet-style decompositions, LSTM models, and gradient-boosted regressors."
    ))

    add_heading(doc, "VIII. Conclusion", 1)
    add_paragraph(doc, (
        "This paper presented an IEEE-style study of an AI-powered smartphone market prediction system that converts public review "
        "text into next-month rating forecasts. The system integrates scraping, cleaning, sentiment classification, ensemble "
        "learning, and hybrid time-series forecasting. Results on 4,021 reviews show that no single sentiment model dominates "
        "all criteria: BERT provides the highest baseline accuracy, soft voting gives the strongest balanced classification and "
        "rating correlation, and bagging minimizes RMSE. By combining sentiment with linear and exponential trend signals, the "
        "system generates realistic product-level forecasts and demonstrates the practical value of sentiment-aware ensemble "
        "learning for market intelligence."
    ))

    add_heading(doc, "References", 1)
    refs = [
        "B. Liu, Sentiment Analysis and Opinion Mining. San Rafael, CA, USA: Morgan & Claypool, 2012.",
        "W. Medhat, A. Hassan, and H. Korashy, \"Sentiment analysis algorithms and applications: A survey,\" Ain Shams Engineering Journal, vol. 5, no. 4, pp. 1093-1113, 2014.",
        "B. Pang, L. Lee, and S. Vaithyanathan, \"Thumbs up? Sentiment classification using machine learning techniques,\" in Proc. Conf. Empirical Methods in Natural Language Processing, 2002, pp. 79-86.",
        "P. D. Turney, \"Thumbs up or thumbs down? Semantic orientation applied to unsupervised classification of reviews,\" in Proc. 40th Annu. Meeting Assoc. Comput. Linguistics, 2002, pp. 417-424.",
        "M. Hu and B. Liu, \"Mining and summarizing customer reviews,\" in Proc. 10th ACM SIGKDD Int. Conf. Knowledge Discovery and Data Mining, 2004, pp. 168-177.",
        "C. J. Hutto and E. Gilbert, \"VADER: A parsimonious rule-based model for sentiment analysis of social media text,\" in Proc. Int. AAAI Conf. Web and Social Media, vol. 8, no. 1, 2014, pp. 216-225.",
        "A. Vaswani et al., \"Attention is all you need,\" in Advances in Neural Information Processing Systems, vol. 30, 2017, pp. 5998-6008.",
        "J. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, \"BERT: Pre-training of deep bidirectional transformers for language understanding,\" in Proc. NAACL-HLT, 2019, pp. 4171-4186.",
        "E. S. Gardner, Jr., \"Exponential smoothing: The state of the art,\" Journal of Forecasting, vol. 4, no. 1, pp. 1-28, 1985.",
        "L. Breiman, \"Bagging predictors,\" Machine Learning, vol. 24, no. 2, pp. 123-140, 1996.",
        "J. H. Friedman, \"Greedy function approximation: A gradient boosting machine,\" Annals of Statistics, vol. 29, no. 5, pp. 1189-1232, 2001.",
        "S. Hochreiter and J. Schmidhuber, \"Long short-term memory,\" Neural Computation, vol. 9, no. 8, pp. 1735-1780, 1997.",
        "A. S. Weigend, Time Series Prediction: Forecasting the Future and Understanding the Past. Reading, MA, USA: Addison-Wesley, 1994.",
        "T. G. Dietterich, \"Ensemble methods in machine learning,\" in Multiple Classifier Systems, Berlin, Germany: Springer, 2000, pp. 1-15.",
        "M. Taboada, J. Brooke, M. Tofiloski, K. Voll, and M. Stede, \"Lexicon-based methods for sentiment analysis,\" Computational Linguistics, vol. 37, no. 2, pp. 267-307, 2011.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.add_run(f"[{i}] {ref}")

    doc.core_properties.title = "AI-Powered Predictive Analysis of Smartphone Market Trends"
    doc.core_properties.subject = "IEEE-style research paper"
    doc.core_properties.author = "Yash Singhal; Yash Khandelwal; Megha Kundi"
    doc.core_properties.keywords = "sentiment analysis, smartphone reviews, ensemble learning, forecasting"
    doc.save(OUT)


if __name__ == "__main__":
    main()
