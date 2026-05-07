import os
from docx import Document
from docx.shared import Pt, RGBColor

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.name = 'Arial'
    if level == 1:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)

def add_math(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(36)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0, 102, 0)

def main():
    doc = Document()
    doc.add_heading('Mathematical Implementation Overview', 0)

    # Slide 1
    add_heading(doc, 'Slide 1: VADER Sentiment Analysis (Lexicon-Based)', 1)
    add_heading(doc, 'Concept', 2)
    add_bullet(doc, 'Rule-based sentiment model using predefined word scores')
    add_bullet(doc, 'Considers word polarity, intensifiers, punctuation, and conjunctions')
    
    add_heading(doc, 'Mathematical Model', 2)
    add_math(doc, 'Compound Score = S / sqrt(S^2 + alpha)')
    add_math(doc, 'Rating = (Compound × 2) + 3')
    
    add_heading(doc, 'Example', 2)
    add_bullet(doc, 'Review: "I absolutely love the camera, it\'s amazing!"')
    add_math(doc, 'Raw Score = 6.825')
    add_math(doc, 'Compound = 0.87')
    add_math(doc, 'Final Rating = (0.87 × 2) + 3 = 4.74')

    # Slide 2
    add_heading(doc, 'Slide 2: Word Cloud Model (Custom Lexicon)', 1)
    add_heading(doc, 'Concept', 2)
    add_bullet(doc, 'Simple keyword-based scoring')
    add_bullet(doc, 'Domain-specific sentiment tracking')
    
    add_heading(doc, 'Mathematical Model', 2)
    add_math(doc, 'Keyword Ratio = (P - N) / (P + N)')
    add_math(doc, 'Score = (Keyword Ratio × 2) + 3')
    
    add_heading(doc, 'Example', 2)
    add_bullet(doc, 'Review: "love the camera, amazing"')
    add_bullet(doc, 'Positive words (P) = 2, Negative words (N) = 0')
    add_math(doc, 'Keyword Ratio = (2 - 0) / (2 + 0) = 1.0')
    add_math(doc, 'Score = (1.0 × 2) + 3 = 5.0')
    
    # Slide 3
    add_heading(doc, 'Slide 3: BERT Sentiment Model (Deep Learning)', 1)
    add_heading(doc, 'Concept', 2)
    add_bullet(doc, 'Transformer-based contextual NLP model')
    add_bullet(doc, 'Outputs probabilities instead of direct score')
    
    add_heading(doc, 'Mathematical Model', 2)
    add_math(doc, 'Rating = 1.0 + (P_pos × 4.0)')
    add_math(doc, 'Where P_pos is the probability of the sentiment being Positive.')
    
    add_heading(doc, 'Example', 2)
    add_bullet(doc, 'Probabilities: Positive (P_pos) = 0.98')
    add_math(doc, 'Rating = 1.0 + (0.98 × 4.0) = 4.92')

    # Slide 4
    add_heading(doc, 'Slide 4: Ensemble – Hard Voting', 1)
    add_heading(doc, 'Concept', 2)
    add_bullet(doc, 'Majority voting among models, converting scores to discrete classes.')
    add_bullet(doc, 'Takes majority vote (ties broken by VADER).')
    add_bullet(doc, 'Positive → 4.5, Neutral → 3.0, Negative → 1.5')
    
    add_heading(doc, 'Example', 2)
    add_bullet(doc, 'VADER: Positive, BERT: Neutral, Word Cloud: Positive')
    add_math(doc, 'Majority Vote = Positive (Mapped to Rating = 4.5)')

    # Slide 5
    add_heading(doc, 'Slide 5: Ensemble – Soft Voting (Chosen Model)', 1)
    add_heading(doc, 'Concept', 2)
    add_bullet(doc, 'Weighted average of predictive scores based on model reliability')
    add_bullet(doc, 'Preserves full continuous sentiment information')
    
    add_heading(doc, 'Mathematical Model', 2)
    add_math(doc, 'Soft Vote = (0.40 × VADER) + (0.40 × BERT) + (0.20 × Word Cloud)')
    
    add_heading(doc, 'Example', 2)
    add_math(doc, 'Soft Vote = (0.40 × 4.74) + (0.40 × 4.92) + (0.20 × 5.0) = 4.864')

    # Slide 6
    add_heading(doc, 'Slide 6: Ensemble – Bagging (Bootstrap Aggregation)', 1)
    add_heading(doc, 'Concept', 2)
    add_bullet(doc, 'Reduce variance using multiple datasets via sampling with replacement')
    
    add_heading(doc, 'Mathematical Model', 2)
    add_math(doc, 'Final Prediction = Average of predictions from 50 bootstrap samples')
    
    add_heading(doc, 'Example', 2)
    add_bullet(doc, 'Predictions from samples: 4.43, 4.41, 4.45…')
    add_math(doc, 'Final = 4.435')

    # Slide 7
    add_heading(doc, 'Slide 7: Final Forecast (Linear + EMA + Ensemble)', 1)
    add_heading(doc, 'Linear Trend', 2)
    add_math(doc, 'y = mx + b')
    add_math(doc, 'Example: (0.015 × 14) + 4.32 = 4.53')
    
    add_heading(doc, 'Exponential Moving Average (EMA)', 2)
    add_math(doc, 'S_t = alpha * y_t + (1 - alpha) * S_{t-1}')
    add_math(doc, 'Example: 4.615')
    
    add_heading(doc, 'Ensemble Sentiment Signal', 2)
    add_math(doc, 'Average soft vote = 4.160')
    
    add_heading(doc, 'Final Weighted Prediction', 2)
    add_math(doc, 'Prediction = 0.4(4.53) + 0.35(4.615) + 0.25(4.160) = 4.467')

    output_path = os.path.join('ppt', 'mathematical_implementation.docx')
    doc.save(output_path)
    print(f"Successfully generated {output_path}")

if __name__ == '__main__':
    main()
